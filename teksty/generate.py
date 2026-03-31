#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Generuje dokumenty Word do przeglądu tekstów na stronie Sielska Chata.
Jeden plik .docx per podstrona, tabela 2-kolumnowa: AKTUALNY TEKST | NOWY TEKST.
"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUT = os.path.dirname(os.path.abspath(__file__))

# ── Kolory ──────────────────────────────────────────────────────────────────
COLOR_HEADER_BG = RGBColor(0x1A, 0x2E, 0x40)
COLOR_SECTION_BG = RGBColor(0x2E, 0x75, 0x52)
COLOR_ROW_ALT   = RGBColor(0xF5, 0xF7, 0xF9)
COLOR_WHITE     = RGBColor(0xFF, 0xFF, 0xFF)
COLOR_TEXT      = RGBColor(0x1A, 0x2E, 0x40)
COLOR_MUTED     = RGBColor(0x66, 0x77, 0x88)
COLOR_NOTE_BG   = RGBColor(0xFF, 0xF8, 0xE1)
COLOR_NOTE_TEXT = RGBColor(0x7A, 0x5C, 0x00)
COLOR_CMS_BG    = RGBColor(0xEE, 0xF4, 0xFF)
COLOR_CMS_TEXT  = RGBColor(0x2E, 0x4E, 0x9A)

INSTRUKCJA = ("Instrukcja: W kolumnie NOWY TEKST wpisz zmienioną wersję "
              "(lub zostaw bez zmian). Nie modyfikuj kolumny AKTUALNY TEKST.")

# ── Helpers ──────────────────────────────────────────────────────────────────
def set_cell_bg(cell, color):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), str(color)); tcPr.append(shd)

def set_cell_border(cell, sides=('top','bottom','left','right'), color='C8D0D8', sz='4'):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in sides:
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'), 'single'); b.set(qn('w:sz'), sz)
        b.set(qn('w:space'), '0'); b.set(qn('w:color'), color)
        tcBorders.append(b)
    tcPr.append(tcBorders)

def set_cell_padding(cell, top=100, bottom=100, left=150, right=150):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for side, val in [('top',top),('bottom',bottom),('left',left),('right',right)]:
        m = OxmlElement(f'w:{side}')
        m.set(qn('w:w'), str(val)); m.set(qn('w:type'), 'dxa'); tcMar.append(m)
    tcPr.append(tcMar)

def add_label_para(cell, text):
    p = cell.paragraphs[0]; p.clear()
    r = p.add_run(text); r.font.size = Pt(7); r.font.bold = True
    r.font.color.rgb = COLOR_MUTED; r.font.name = 'Calibri'
    p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(2)

def add_content_para(cell, text, bold=False, size=10):
    p = cell.add_paragraph()
    r = p.add_run(text); r.font.size = Pt(size); r.font.bold = bold
    r.font.color.rgb = COLOR_TEXT; r.font.name = 'Calibri'
    p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(0)

def style_header_row(table):
    row = table.rows[0]
    for i, label in enumerate(["AKTUALNY TEKST", "NOWY TEKST"]):
        cell = row.cells[i]
        set_cell_bg(cell, COLOR_HEADER_BG)
        set_cell_border(cell, color='1A2E40', sz='6')
        set_cell_padding(cell, top=120, bottom=120)
        p = cell.paragraphs[0]; p.clear()
        r = p.add_run(label); r.font.size = Pt(9); r.font.bold = True
        r.font.color.rgb = COLOR_WHITE; r.font.name = 'Calibri'
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

def new_doc(title):
    doc = Document()
    s = doc.sections[0]
    s.page_width = Cm(29.7); s.page_height = Cm(21.0)
    s.left_margin = s.right_margin = Cm(1.5)
    s.top_margin = s.bottom_margin = Cm(1.5)
    tp = doc.add_paragraph()
    r = tp.add_run(f"Sielska Chata — Przegląd tekstów: {title}")
    r.font.size = Pt(14); r.font.bold = True
    r.font.color.rgb = COLOR_HEADER_BG; r.font.name = 'Calibri'
    tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tp.paragraph_format.space_after = Pt(4)
    sp = doc.add_paragraph()
    sr = sp.add_run(INSTRUKCJA)
    sr.font.size = Pt(9); sr.font.italic = True
    sr.font.color.rgb = COLOR_MUTED; sr.font.name = 'Calibri'
    sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sp.paragraph_format.space_after = Pt(10)
    t = doc.add_table(rows=1, cols=2)
    t.style = 'Table Grid'; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for col in t.columns:
        for cell in col.cells:
            cell.width = Cm(12.8)
    style_header_row(t)
    return doc, t

def sec(table, title):
    tr = table.add_row(); tr.cells[0].merge(tr.cells[1])
    cell = tr.cells[0]
    set_cell_bg(cell, COLOR_SECTION_BG)
    set_cell_border(cell, color='2E7552', sz='8')
    set_cell_padding(cell, top=130, bottom=130, left=200, right=200)
    p = cell.paragraphs[0]; p.clear()
    r = p.add_run(title); r.font.size = Pt(10); r.font.bold = True
    r.font.color.rgb = COLOR_WHITE; r.font.name = 'Calibri'
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

def row(table, label, text, alt=False, bold=False, size=10):
    tr = table.add_row()
    bg = COLOR_ROW_ALT if alt else COLOR_WHITE
    for ci in range(2):
        cell = tr.cells[ci]
        set_cell_bg(cell, bg)
        set_cell_border(cell, color='D0D8E0', sz='4')
        set_cell_padding(cell, top=100, bottom=100, left=160, right=160)
        add_label_para(cell, label)
        if ci == 0:
            add_content_para(cell, text, bold=bold, size=size)
        else:
            add_content_para(cell, '', bold=bold, size=size)

def note(table, text):
    tr = table.add_row(); tr.cells[0].merge(tr.cells[1])
    cell = tr.cells[0]
    set_cell_bg(cell, COLOR_NOTE_BG)
    set_cell_border(cell, color='E0C060', sz='4')
    set_cell_padding(cell, top=80, bottom=80, left=200, right=200)
    p = cell.paragraphs[0]; p.clear()
    r = p.add_run(f"UWAGA: {text}"); r.font.size = Pt(8); r.font.italic = True
    r.font.color.rgb = COLOR_NOTE_TEXT; r.font.name = 'Calibri'

def cms_note(table, text):
    tr = table.add_row(); tr.cells[0].merge(tr.cells[1])
    cell = tr.cells[0]
    set_cell_bg(cell, COLOR_CMS_BG)
    set_cell_border(cell, color='8AAAE0', sz='4')
    set_cell_padding(cell, top=80, bottom=80, left=200, right=200)
    p = cell.paragraphs[0]; p.clear()
    r = p.add_run(f"CMS: {text}"); r.font.size = Pt(8); r.font.italic = True
    r.font.color.rgb = COLOR_CMS_TEXT; r.font.name = 'Calibri'

def save(doc, filename):
    path = os.path.join(OUT, filename)
    doc.save(path)
    print(f"  {filename}")


# ════════════════════════════════════════════════════════════════════════════
# 1. HOME.DOCX — Strona główna
# ════════════════════════════════════════════════════════════════════════════
doc, t = new_doc("Strona główna")

sec(t, "1. SEKCJA HERO")
row(t, "BADGE / WYRÓŻNIENIE", "Laureat Orłów Gastronomii 2026", alt=False)
row(t, "NAGŁÓWEK H1", "Sielska Chata — Restauracja w Rabce-Zdrój", alt=True, bold=True, size=11)
row(t, "PODTYTUŁ", "Dobra kuchnia. Coś dla każdego.", alt=False)
row(t, "PRZYCISK 1", "Zarezerwuj stolik", alt=True)
row(t, "PRZYCISK 2", "Zobacz menu", alt=False)

sec(t, "2. PASEK WYRÓŻNIEŃ (Awards Banner)")
row(t, "OCENA GOOGLE — wartość", "4.7", alt=False)
row(t, "OCENA GOOGLE — platforma", "Google Maps", alt=True)
row(t, "OCENA GOOGLE — opis", "76 opinii", alt=False)
row(t, "OCENA ORŁY — wartość", "8.9", alt=True)
row(t, "OCENA ORŁY — platforma", "Orły Gastronomii", alt=False)
row(t, "OCENA ORŁY — opis", "Laureat 2026", alt=True)
row(t, "GODZINY — wartość", "9:00–19:00", alt=False)
row(t, "GODZINY — opis 1", "Codziennie", alt=True)
row(t, "GODZINY — opis 2", "7 dni w tygodniu", alt=False)
row(t, "PARKING — wartość", "Duży", alt=True)
row(t, "PARKING — opis 1", "Parking", alt=False)
row(t, "PARKING — opis 2", "Bezpłatny", alt=True)

sec(t, "3. SEKCJA O NAS")
row(t, "LABEL (mały tekst nad nagłówkiem)", "O nas", alt=False)
row(t, "NAGŁÓWEK H2", "Rodzinna restauracja, drewniane wnętrze i uczciwa kuchnia z długą listą pyszności.", alt=True, bold=True, size=11)
row(t, "TEKST AKAPIT 1", "Nasze obszerne menu to nasz największy atut. Dbamy o to, by łącząc pokolenia przy jednym stole, zadowolić każdy gust: od ulubionych przez dzieci klopsików i frytek, przez wykwintną jagnięciną dla rodziców, aż po domowy rosół dla dziadków. Każdy gość wychodzi od nas najedzony i uśmiechnięty.", alt=False)
row(t, "TEKST AKAPIT 2", "Wkładamy w gotowanie całe serce, traktując gości jak bliskich. Zrób sobie przerwę w podróży pod Tatry – mieścimy się przy trasie Kraków–Zakopane, tylko 5 minut od Parku Zdrojowego w Rabce-Zdroju.", alt=True)
row(t, "LINK", "Więcej o nas", alt=False)

sec(t, "4. KARTY CECH (3 karty)")
row(t, "KARTA 1 — TYTUŁ", "Polska klasyka i europejskie smaki", alt=False, bold=True)
row(t, "KARTA 1 — OPIS", "Od pierogów, golonki i żurku, przez żeberka BBQ i pstrąga, aż po jagnięcinę, sandacza w borowikach i krewetki. Coś dobrego dla każdego.", alt=True)
row(t, "KARTA 2 — TYTUŁ", "Imprezy okolicznościowe", alt=False, bold=True)
row(t, "KARTA 2 — OPIS", "Komunia, chrzciny, urodziny, bankiet firmowy — organizujemy od A do Z. Menu dostosujemy do potrzeb i liczby gości.", alt=True)
row(t, "KARTA 3 — TYTUŁ", "W sercu uzdrowiska", alt=False, bold=True)
row(t, "KARTA 3 — OPIS", "Centrum Rabki-Zdrój, trasa Kraków–Zakopane. Park Zdrojowy, Rabkoland, Luboń Wielki — tuż obok.", alt=True)

sec(t, "5. PODGLĄD MENU (Karta dań)")
row(t, "LABEL", "Karta dań", alt=False)
row(t, "NAGŁÓWEK H2", "Kuchnia polska i europejska", alt=True, bold=True, size=11)
row(t, "OPIS", "Tradycyjne polskie smaki, ryby z górskich potoków, europejskie dania mięsne i szeroka karta napojów. Menu zmienia się sezonowo.", alt=False)
row(t, "PRZYCISK", "Zobacz pełne menu", alt=True)
row(t, "DOPISEK", "Dostępne również jedzenie na wynos w Rabce-Zdrój", alt=False)

sec(t, "6. CYTAT / FILOZOFIA")
row(t, "CYTAT", "Nie obiecujemy rewolucji. Obiecujemy solidny obiad, ciepłe wnętrze i spokój, którego szukasz na urlopie.", alt=False)
row(t, "PODTEKST", "Sielska Chata to miejsce, gdzie czas zwalnia. Drewniane wnętrze, szeroka karta, duże porcje. Siadasz, zamawiasz bez pośpiechu, a my podajemy to, co gotujemy od lat — uczciwie.", alt=True)

sec(t, "7. SEKCJA IMPREZY (teaser)")
row(t, "LABEL", "Imprezy okolicznościowe", alt=False)
row(t, "NAGŁÓWEK H2", "Zorganizuj imprezę w Rabce-Zdrój", alt=True, bold=True, size=11)
row(t, "OPIS", "Komunia, chrzciny, urodziny, spotkanie firmowe — dostosujemy salę i menu do każdej okazji.", alt=False)
row(t, "KARTA 1 — TYTUŁ", "Komunia i chrzciny", alt=True, bold=True)
row(t, "KARTA 1 — OPIS", "Ciepłe, rodzinne przyjęcie dla kilkunastu lub kilkudziesięciu gości. Menu dostosowane do dzieci i dorosłych.", alt=False)
row(t, "KARTA 2 — TYTUŁ", "Urodziny i jubileusze", alt=True, bold=True)
row(t, "KARTA 2 — OPIS", "Urodziny w drewnianym wnętrzu — ciepło i bez stresu. Menu, dekoracje, liczba gości — wszystko ustalimy.", alt=False)
row(t, "KARTA 3 — TYTUŁ", "Imprezy firmowe i bankiety", alt=True, bold=True)
row(t, "KARTA 3 — OPIS", "Spotkania integracyjne, kolacje firmowe, bankiety. Obsługa do kilkudziesięciu osób.", alt=False)
row(t, "PRZYCISK", "Sprawdź ofertę imprez", alt=True)

sec(t, "8. OPINIE GOŚCI")
row(t, "LABEL", "Opinie gości", alt=False)
row(t, "NAGŁÓWEK H2", "Co mówią o nas", alt=True, bold=True, size=11)
row(t, "INFO OCENY", "Google 4.7/5  |  Orły Gastronomii 8.9/10", alt=False)
row(t, "OPINIA 1 — treść", "Najlepsza kwaśnica jaką jadłem poza domem babci. Porcje ogromne, a obsługa serdeczna jak u rodziny.", alt=True)
row(t, "OPINIA 1 — autor", "Marek K., Turysta z Krakowa", alt=False)
row(t, "OPINIA 2 — treść", "Robiliśmy tu komunię córki — 40 osób, zero stresu. Wszystko dopięte, a goście pytali o przepis na sernik.", alt=True)
row(t, "OPINIA 2 — autor", "Anna W., Organizacja komunii", alt=False)
row(t, "OPINIA 3 — treść", "Wracamy tu co wakacje. Dzieci uwielbiają naleśniki, my — spokój i widok z okna. Parking ogromny.", alt=True)
row(t, "OPINIA 3 — autor", "Tomasz i Kasia, Rodzina na urlopie", alt=False)

sec(t, "9. SEKCJA KONTAKT (teaser)")
row(t, "LABEL", "Kontakt", alt=False)
row(t, "NAGŁÓWEK H2", "Zapraszamy do Sielskiej Chaty w Rabce-Zdrój", alt=True, bold=True, size=11)
row(t, "TEKST", "Szukasz dobrego miejsca, gdzie zjeść w Rabce-Zdrój? Czynni codziennie 9:00–19:00. Zarezerwuj stolik telefonicznie lub wpadnij spontanicznie — prawie zawsze znajdziemy miejsce.", alt=False)
row(t, "PRZYCISK 1", "Zarezerwuj stolik", alt=True)
row(t, "PRZYCISK 2", "Jak dojechać", alt=False)
row(t, "ADRES", "ul. Piłsudskiego 18, 34-700 Rabka-Zdrój", alt=True)
row(t, "TELEFON", "+48 780 285 859", alt=False)
row(t, "GODZINY", "Codziennie 9:00–19:00", alt=True)
row(t, "PARKING", "Duży, bezpłatny parking przy restauracji", alt=False)

save(doc, "home.docx")


# ════════════════════════════════════════════════════════════════════════════
# 2. MENU.DOCX — Menu
# ════════════════════════════════════════════════════════════════════════════
doc, t = new_doc("Menu")

sec(t, "1. SEKCJA HERO")
row(t, "LABEL", "Karta dań", alt=False)
row(t, "NAGŁÓWEK H1", "Menu — Sielska Chata / Rabka-Zdrój", alt=True, bold=True, size=11)
row(t, "OPIS", "Szeroka karta z kuchnią polską i europejską. Przystawki, zupy, dania główne, ryby, pierogi, sałatki, desery i menu dla dzieci. Ceny orientacyjne — menu zmienia się sezonowo.", alt=False)

sec(t, "2. KARTY INFO (3 karty)")
row(t, "KARTA 1 — TYTUŁ", "Jedzenie na wynos", alt=False, bold=True)
row(t, "KARTA 1 — OPIS", "Zamów telefonicznie — odbierz gotowe", alt=True)
row(t, "KARTA 2 — TYTUŁ", "Menu dla dzieci", alt=False, bold=True)
row(t, "KARTA 2 — OPIS", "Rosół, pomidorowa, klopsiki, nuggetsy", alt=True)
row(t, "KARTA 3 — TYTUŁ", "Menu sezonowe", alt=False, bold=True)
row(t, "KARTA 3 — OPIS", "Karta zmienia się przez cały rok", alt=True)

sec(t, "3. SEKCJA DANIE DNIA")
row(t, "NAGŁÓWEK H2", "Danie dnia", alt=False, bold=True, size=11)
row(t, "TEKST", "Danie dnia publikujemy codziennie na naszym profilu na Facebooku. Zapraszamy do śledzenia — tam też najszybciej dowiesz się o aktualnej ofercie i promocjach.", alt=True)
row(t, "PRZYCISK", "Zobacz na Facebooku", alt=False)

sec(t, "4. SEKCJA GALERIA DAŃ")
row(t, "LABEL", "Galeria", alt=False)
row(t, "NAGŁÓWEK H2", "Nasze dania", alt=True, bold=True, size=11)
cms_note(t, "Zdjęcia galerii — klient dostarcza własne zdjęcia dań. Aktualnie: placeholder hero.jpg ×5.")

sec(t, "5. SEKCJA JEDZENIE NA WYNOS")
row(t, "NAGŁÓWEK H2", "Jedzenie na wynos — Rabka-Zdrój", alt=False, bold=True, size=11)
row(t, "TEKST", "Zamów telefonicznie, przyjdź i odbierz. Czynni codziennie 9:00–19:00.", alt=True)
row(t, "PRZYCISK 1", "Zadzwoń i zamów", alt=False)
row(t, "PRZYCISK 2", "Jak dojechać", alt=True)

sec(t, "6. STOPKA MENU")
row(t, "UWAGA / DOPISEK", "Ceny orientacyjne. Menu zmienia się sezonowo.", alt=False)

save(doc, "menu.docx")


# ════════════════════════════════════════════════════════════════════════════
# 3. O-NAS.DOCX — O nas
# ════════════════════════════════════════════════════════════════════════════
doc, t = new_doc("O nas")

sec(t, "1. SEKCJA HERO")
row(t, "LABEL", "O nas", alt=False)
row(t, "NAGŁÓWEK H1", "O nas — Sielska Chata / Rabka-Zdrój", alt=True, bold=True, size=11)
row(t, "OPIS", "Restauracja rodzinna w centrum uzdrowiska. Kuchnia polska i europejska, uczciwe gotowanie, drewniane wnętrze — i karta dla każdego apetytu.", alt=False)

sec(t, "2. TREŚĆ O NAS")
row(t, "NAGŁÓWEK H2", "Rodzinna restauracja, drewniane wnętrze i uczciwa kuchnia z długą listą pyszności.", alt=False, bold=True, size=11)
row(t, "TEKST AKAPIT 1", "Wkładamy w gotowanie całe serce, traktując gości jak bliskich. Zrób sobie przerwę w podróży pod Tatry – mieścimy się przy trasie Kraków–Zakopane, tylko 5 minut od Parku Zdrojowego w Rabce-Zdroju.", alt=True)
row(t, "TEKST AKAPIT 2", "Sielska Chata to restauracja w centrum Rabki-Zdrój. Prowadzimy ją z myślą o gościach, którzy chcą zjeść dobrze — bez poszukiwania wyjątkowych odkryć kulinarnych, ale z pewnością solidnego, uczciwego posiłku.", alt=False)
row(t, "TEKST AKAPIT 3", "Serwujemy kuchnię polską i europejską. Mamy w menu tradycyjne polskie dania jak pierogi, żurek, kwaśnicę, placki po zbójnicku, golonkę pieczoną w piwie. Ale też bardziej europejskie propozycje — ślimaki w maśle czosnkowym, stek wołowy, pstrąga z ziołami, filet z sandacza, krewetki z ryżem. Menu zmienia się sezonowo.", alt=True)
row(t, "TEKST AKAPIT 4 (nagrody)", "Orły Gastronomii 2026 — ocena 8.9. Na Google Maps 4.7 z ponad 76 recenzji. Oba cieszą, ale bardziej liczy się to, co czujesz wychodząc — i czy wróciłbyś w przyszłym tygodniu.", alt=False)

sec(t, "3. KARTY CECH (3 karty)")
row(t, "KARTA 1 — TYTUŁ", "Kuchnia polska i europejska", alt=False, bold=True)
row(t, "KARTA 1 — OPIS", "Tradycyjne polskie smaki — pierogi, żurek, golonka, kwaśnica. Ale też ślimaki, jagnięcina, krewetki i ryby z górskich potoków.", alt=True)
row(t, "KARTA 2 — TYTUŁ", "Laureat Orłów Gastronomii 2026", alt=False, bold=True)
row(t, "KARTA 2 — OPIS", "Ocena 8.9/10 w prestiżowym rankingu gastronomicznym. Google Maps 4.7/5 — 76 opinii gości.", alt=True)
row(t, "KARTA 3 — TYTUŁ", "Centrum Rabki-Zdrój", alt=False, bold=True)
row(t, "KARTA 3 — OPIS", "Ul. Piłsudskiego 18, trasa Kraków–Zakopane. Park Zdrojowy, tężnia, Rabkoland — wszystko tuż obok.", alt=True)

sec(t, "4. CYTAT")
row(t, "CYTAT", "Nie obiecujemy rewolucji. Obiecujemy solidny obiad, ciepłe wnętrze i spokój, którego szukasz na urlopie.", alt=False)
row(t, "ATRYBUCJA", "Sielska Chata, Rabka-Zdrój", alt=True)

sec(t, "5. SEKCJA CTA")
row(t, "NAGŁÓWEK H2", "Zapraszamy do Sielskiej Chaty", alt=False, bold=True, size=11)
row(t, "PRZYCISK 1", "Zadzwoń do nas", alt=True)
row(t, "PRZYCISK 2", "Zobacz menu", alt=False)

save(doc, "o-nas.docx")


# ════════════════════════════════════════════════════════════════════════════
# 4. KONTAKT.DOCX — Kontakt
# ════════════════════════════════════════════════════════════════════════════
doc, t = new_doc("Kontakt")

sec(t, "1. SEKCJA HERO")
row(t, "LABEL", "Kontakt", alt=False)
row(t, "NAGŁÓWEK H1", "Kontakt i Rezerwacja / Sielska Chata Rabka-Zdrój", alt=True, bold=True, size=11)
row(t, "OPIS", "Zarezerwuj stolik telefonicznie lub wpadnij spontanicznie. Czynni codziennie 9:00–19:00. Bezpłatny parking przy restauracji.", alt=False)

sec(t, "2. DANE KONTAKTOWE")
row(t, "NAGŁÓWEK H2", "Zapraszamy", alt=False, bold=True, size=11)
row(t, "TEKST INTRO", "Masz pytanie o menu, chcesz zarezerwować stolik albo porozmawiać o organizacji imprezy? Zadzwoń — odpowiadamy szybko.", alt=True)
row(t, "ADRES", "ul. Piłsudskiego 18, 34-700 Rabka-Zdrój", alt=False)
row(t, "TELEFON", "+48 780 285 859", alt=True)
row(t, "GODZINY OTWARCIA", "Codziennie 9:00–19:00", alt=False)
row(t, "PARKING", "Duży, bezpłatny parking przy restauracji", alt=True)
row(t, "PRZYCISK", "Zadzwoń teraz", alt=False)

sec(t, "3. SEKCJA JAK DOJECHAĆ")
row(t, "NAGŁÓWEK H2", "Jak dojechać do Sielskiej Chaty?", alt=False, bold=True, size=11)
row(t, "KIERUNEK 1 — TYTUŁ", "Samochodem z Krakowa", alt=True, bold=True)
row(t, "KIERUNEK 1 — OPIS", "Drogą krajową DK47 przez Myślenice i Rabkę. Rabka-Zdrój to ok. 70 km od centrum Krakowa — około godziny jazdy. Restauracja przy ul. Piłsudskiego 18, tuż przy głównej drodze.", alt=False)
row(t, "KIERUNEK 2 — TYTUŁ", "Samochodem z Zakopanego", alt=True, bold=True)
row(t, "KIERUNEK 2 — OPIS", "Drogą DK47 w kierunku Rabki-Zdrój — ok. 25 km, 30 minut. Sielska Chata to świetny przystanek w drodze powrotnej z Zakopanego.", alt=False)
row(t, "KIERUNEK 3 — TYTUŁ", "Komunikacją publiczną", alt=True, bold=True)
row(t, "KIERUNEK 3 — OPIS", "Pociągiem lub autobusem do Rabki-Zdrój. Dworzec autobusowy i PKP w odległości kilku minut spacerem od restauracji.", alt=False)
row(t, "KIERUNEK 4 — TYTUŁ", "Parking", alt=True, bold=True)
row(t, "KIERUNEK 4 — OPIS", "Duży, bezpłatny parking bezpośrednio przy restauracji. Nie musisz szukać miejsca — zawsze znajdziesz wolne.", alt=False)

save(doc, "kontakt.docx")


# ════════════════════════════════════════════════════════════════════════════
# 5. IMPREZY.DOCX — Imprezy okolicznościowe
# ════════════════════════════════════════════════════════════════════════════
doc, t = new_doc("Imprezy okolicznościowe")

sec(t, "1. SEKCJA HERO")
row(t, "LABEL", "Imprezy okolicznościowe", alt=False)
row(t, "NAGŁÓWEK H1", "Imprezy Okolicznościowe / w Rabce-Zdrój", alt=True, bold=True, size=11)
row(t, "OPIS", "Komunia, chrzciny, urodziny, jubileusze, bankiety firmowe — organizujemy uroczystości od A do Z w sercu Rabki-Zdrój. Klimatyczne wnętrze, kuchnia polska i europejska, bezpłatny parking.", alt=False)

sec(t, "2. QUICK LINKS (nawigacja na stronie)")
row(t, "LINK 1", "Komunia Święta", alt=False)
row(t, "LINK 2", "Chrzciny i Przyjęcia Chrzcielne", alt=True)
row(t, "LINK 3", "Urodziny i Jubileusze", alt=False)
row(t, "LINK 4", "Imprezy Firmowe i Bankiety", alt=True)

sec(t, "3. KARTA 01 — KOMUNIA ŚWIĘTA")
row(t, "NUMER", "01", alt=False)
row(t, "NAGŁÓWEK H2", "Komunia Święta", alt=True, bold=True, size=11)
row(t, "PODTYTUŁ", "Organizacja komunii w Rabce-Zdrój", alt=False)
row(t, "OPIS", "Komunia to dzień, po którym wspomnienia zostają na lata. W Sielskiej Chacie w Rabce-Zdrój dbamy o to, żeby część przy stole była bez kłopotów — menu dopasowane do dzieci i dorosłych, klimatyczna sala z drewnianym wystrojem, obsługa, która zna takie uroczystości. Ty skupiasz się na dziecku.", alt=True)
row(t, "CECHA 1", "Menu dla dzieci i dorosłych — każdy znajdzie coś swojego", alt=False)
row(t, "CECHA 2", "Klimatyczna sala z drewnianym wystrojem", alt=True)
row(t, "CECHA 3", "Obsługa grup od kilkunastu do kilkudziesięciu osób", alt=False)
row(t, "CECHA 4", "Możliwość dekoracji sali na tę okazję", alt=True)
row(t, "CECHA 5", "Bezpłatny parking przy restauracji", alt=False)
row(t, "OPINIA — treść", "Robiliśmy tu komunię córki — 40 osób, zero stresu. Wszystko dopięte na ostatni guzik.", alt=True)
row(t, "OPINIA — autor", "Anna W., organizacja komunii", alt=False)

sec(t, "4. KARTA 02 — CHRZCINY I PRZYJĘCIA CHRZCIELNE")
row(t, "NUMER", "02", alt=False)
row(t, "NAGŁÓWEK H2", "Chrzciny i Przyjęcia Chrzcielne", alt=True, bold=True, size=11)
row(t, "PODTYTUŁ", "Chrzciny w restauracji Rabka-Zdrój", alt=False)
row(t, "OPIS", "Po ceremonii w kościele wszyscy chcą po prostu usiąść i zjeść razem. Organizujemy chrzciny w drewnianym, ciepłym wnętrzu Sielskiej Chaty — centrum Rabki-Zdrój, kilka minut od głównych kościołów. Elastyczne menu, spokojna atmosfera, parking przy wejściu.", alt=True)
row(t, "CECHA 1", "Menu do wyboru — kuchnia polska i europejska", alt=False)
row(t, "CECHA 2", "Centrum Rabki-Zdrój — blisko kościołów i Parku Zdrojowego", alt=True)
row(t, "CECHA 3", "Możliwość zamówienia tortu lub ciasta na uroczystość", alt=False)
row(t, "CECHA 4", "Obsługa grup rodzinnych od kilkunastu osób", alt=True)
row(t, "CECHA 5", "Duży bezpłatny parking", alt=False)

sec(t, "5. KARTA 03 — URODZINY I JUBILEUSZE")
row(t, "NUMER", "03", alt=False)
row(t, "NAGŁÓWEK H2", "Urodziny i Jubileusze", alt=True, bold=True, size=11)
row(t, "PODTYTUŁ", "Urodziny w restauracji Rabka-Zdrój", alt=False)
row(t, "OPIS", "Urodziny w restauracji Rabka-Zdrój — bez planowania na ostatnią chwilę. Sielska Chata to klimatyczne drewniane wnętrze, karta, z której każdy coś wybierze, i obsługa, która nie potrzebuje instrukcji. Ty skupiasz się na gościach, my na reszcie.", alt=True)
row(t, "CECHA 1", "Rezerwacja sali lub wydzielonej przestrzeni", alt=False)
row(t, "CECHA 2", "Indywidualne menu na zamówienie", alt=True)
row(t, "CECHA 3", "Możliwość wniesienia własnego tortu", alt=False)
row(t, "CECHA 4", "Kuchnia polska i europejska — szeroki wybór dla każdego", alt=True)
row(t, "CECHA 5", "Czynni codziennie 9:00–19:00", alt=False)

sec(t, "6. KARTA 04 — IMPREZY FIRMOWE I BANKIETY")
row(t, "NUMER", "04", alt=False)
row(t, "NAGŁÓWEK H2", "Imprezy Firmowe i Bankiety", alt=True, bold=True, size=11)
row(t, "PODTYTUŁ", "Bankiet i spotkanie firmowe w Rabce-Zdrój", alt=False)
row(t, "OPIS", "Szukasz miejsca na bankiet firmowy lub spotkanie integracyjne w okolicach Rabki-Zdrój? Sielska Chata to dobra lokalizacja przy trasie Kraków–Zakopane, duży parking, przestrzeń dla kilkudziesięciu osób i kuchnia polska z europejskim zacięciem. Menu ustalamy pod konkretną grupę.", alt=True)
row(t, "CECHA 1", "Doświadczenie w obsłudze grup od 15 do 60 osób", alt=False)
row(t, "CECHA 2", "Menu dostosowane do preferencji i budżetu", alt=True)
row(t, "CECHA 3", "Kuchnia polska i europejska — dobre na mieszane grupy", alt=False)
row(t, "CECHA 4", "Lokalizacja przy trasie Kraków–Zakopane, duży parking", alt=True)
row(t, "CECHA 5", "Możliwość rezerwacji całej sali", alt=False)

sec(t, "7. FAQ — JAK ZAREZERWOWAĆ")
row(t, "LABEL", "Jak zarezerwować", alt=False)
row(t, "NAGŁÓWEK H2", "Prosto i bez komplikacji", alt=True, bold=True, size=11)
row(t, "PYTANIE 1", "Jak zarezerwować salę na imprezę w Sielskiej Chacie?", alt=False, bold=True)
row(t, "ODPOWIEDŹ 1", "Najprościej zadzwoń pod numer +48 780 285 859. Omówimy datę, liczbę gości i rodzaj imprezy. Możesz też wpaść osobiście — jesteśmy czynni codziennie 9:00–19:00.", alt=True)
row(t, "PYTANIE 2", "Na ile osób można zorganizować przyjęcie?", alt=False, bold=True)
row(t, "ODPOWIEDŹ 2", "Przyjmujemy grupy od kilkunastu do kilkudziesięciu osób. Skontaktuj się z nami, żeby ustalić szczegóły dotyczące Twojej uroczystości.", alt=True)
row(t, "PYTANIE 3", "Czy menu na imprezę można dostosować?", alt=False, bold=True)
row(t, "ODPOWIEDŹ 3", "Oczywiście. Nasza szeroka karta pozwala skomponować menu pasujące do każdej okazji — od tradycyjnych polskich dań, przez ryby i dania europejskie, aż po menu dla dzieci.", alt=True)
row(t, "PYTANIE 4", "Czy jest parking przy restauracji?", alt=False, bold=True)
row(t, "ODPOWIEDŹ 4", "Tak — dysponujemy dużym, bezpłatnym parkingiem. Goście przyjeżdżający z całej Małopolski nie mają problemów z miejscem.", alt=True)

sec(t, "8. SEKCJA CTA")
row(t, "NAGŁÓWEK H2", "Zaplanuj imprezę w Rabce-Zdrój", alt=False, bold=True, size=11)
row(t, "TEKST", "Zadzwoń do nas — ustalimy datę, menu i wszystkie szczegóły Twojej uroczystości.", alt=True)
row(t, "PRZYCISK", "+48 780 285 859", alt=False)

save(doc, "imprezy.docx")

print("\nGotowe! 5 dokumentow Word w folderze teksty/")
