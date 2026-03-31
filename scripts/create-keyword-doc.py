#!/usr/bin/env python3
"""Generates SEO keyword strategy document for Sielska Chata."""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

doc = Document()

# ── Styles ────────────────────────────────────────────────────────────────────

def set_font(run, size=11, bold=False, color=None, italic=False):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.name = "Calibri"
    if color:
        run.font.color.rgb = RGBColor(*color)

def add_heading(doc, text, level=1, color=(44, 36, 24)):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in p.runs:
        run.font.name = "Calibri"
        run.font.color.rgb = RGBColor(*color)
    return p

def add_table_row(table, cells_data, bold_first=False, header=False):
    row = table.add_row()
    for i, (cell, data) in enumerate(zip(row.cells, cells_data)):
        p = cell.paragraphs[0]
        run = p.add_run(data)
        run.font.name = "Calibri"
        run.font.size = Pt(10)
        if header or (bold_first and i == 0):
            run.font.bold = True
        if header:
            run.font.color.rgb = RGBColor(255, 255, 255)
            cell._tc.get_or_add_tcPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:fill'), '2C2418')
            shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:val'), 'clear')
            cell._tc.tcPr.append(shd)
    return row

# ── Page margins ──────────────────────────────────────────────────────────────
sections = doc.sections
for section in sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# ── Title page ────────────────────────────────────────────────────────────────
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title_p.add_run("STRATEGIA SEO")
set_font(run, size=28, bold=True, color=(44, 36, 24))

subtitle_p = doc.add_paragraph()
subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle_p.add_run("Sielska Chata — Restauracja Rabka-Zdrój")
set_font(run, size=16, color=(180, 100, 60))

doc.add_paragraph()
date_p = doc.add_paragraph()
date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = date_p.add_run("Marzec 2026 · aisolutions.design")
set_font(run, size=10, color=(120, 110, 100))

doc.add_page_break()

# ── 1. PODSUMOWANIE ───────────────────────────────────────────────────────────
add_heading(doc, "1. Podsumowanie strategii", level=1)

p = doc.add_paragraph()
run = p.add_run(
    "Celem strategii SEO dla Sielskiej Chaty jest osiągnięcie widoczności w Google dla lokalnych zapytań "
    "związanych z restauracją w Rabce-Zdrój. Restauracja posiada silną przewagę konkurencyjną: "
    "Laureat Orłów Gastronomii 2026 (8.9/10), ocena Google 4.7/5, szeroka karta kuchni polskiej i europejskiej, "
    "imprezy okolicznościowe i jedzenie na wynos."
)
set_font(run)

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run("Priorytet: ")
set_font(run, bold=True)
run2 = p.add_run(
    "pozycjonowanie lokalne — frazy z geolokalizacją 'Rabka-Zdrój' i 'Rabka'. "
    "Horyzont: top 5 Google dla głównych fraz w ciągu 3–6 miesięcy od wdrożenia."
)
set_font(run2)

# ── 2. ARCHITEKTURA STRON ────────────────────────────────────────────────────
doc.add_paragraph()
add_heading(doc, "2. Architektura podstron", level=1)

p = doc.add_paragraph()
run = p.add_run("Strona podzielona na 5 indeksowanych URL z unikalnymi treściami SEO:")
set_font(run)

table = doc.add_table(rows=1, cols=3)
table.style = "Table Grid"
table.alignment = WD_TABLE_ALIGNMENT.LEFT

# Header
add_table_row(table, ["URL", "Tytuł strony", "Główne frazy"], header=True)

pages = [
    ("/", "Sielska Chata — Restauracja w Rabce-Zdrój", "restauracja Rabka-Zdrój, gdzie zjeść Rabka"),
    ("/menu/", "Menu — Sielska Chata Rabka-Zdrój | Kuchnia polska i europejska", "menu restauracja Rabka, obiad Rabka-Zdrój"),
    ("/imprezy/", "Imprezy Okolicznościowe Rabka-Zdrój | Komunia, Chrzciny, Urodziny", "imprezy Rabka, komunia Rabka, chrzciny Rabka"),
    ("/o-nas/", "O nas — Sielska Chata | Restauracja w centrum Rabki-Zdrój", "Sielska Chata Rabka, o restauracji"),
    ("/kontakt/", "Kontakt i Rezerwacja — Sielska Chata Rabka-Zdrój", "kontakt restauracja Rabka, rezerwacja stolika"),
]

for page in pages:
    add_table_row(table, page, bold_first=True)

# ── 3. SŁOWA KLUCZOWE ─────────────────────────────────────────────────────────
doc.add_paragraph()
add_heading(doc, "3. Słowa kluczowe", level=1)

# 3.1 PRIMARY
add_heading(doc, "3.1 Podstawowe (strona główna)", level=2)
p = doc.add_paragraph()
run = p.add_run("Najwyższy ruch, największa konkurencja — kluczowe dla homepage.")
set_font(run, italic=True, color=(100, 90, 80))

table2 = doc.add_table(rows=1, cols=4)
table2.style = "Table Grid"
add_table_row(table2, ["Fraza", "Szac. wolumen/mc", "Trudność", "Docelowa strona"], header=True)

primary = [
    ("restauracja Rabka-Zdrój", "500–1000", "Średnia", "/"),
    ("restauracja Rabka", "300–700", "Średnia", "/"),
    ("kuchnia polska Rabka-Zdrój", "100–200", "Niska-Średnia", "/"),
    ("obiad Rabka-Zdrój", "100–200", "Niska", "/"),
    ("gdzie zjeść Rabka-Zdrój", "200–400", "Niska", "/"),
    ("jedzenie na wynos Rabka-Zdrój", "100–200", "Niska-Średnia", "/"),
    ("najlepsza restauracja Rabka-Zdrój", "50–100", "Niska", "/"),
    ("restauracja dla dzieci Rabka-Zdrój", "50–100", "Niska", "/"),
]

for row in primary:
    add_table_row(table2, row)

doc.add_paragraph()

# 3.2 SECONDARY — imprezy
add_heading(doc, "3.2 Imprezy okolicznościowe (/imprezy/)", level=2)

table3 = doc.add_table(rows=1, cols=4)
table3.style = "Table Grid"
add_table_row(table3, ["Fraza", "Szac. wolumen/mc", "Trudność", "Docelowa strona"], header=True)

secondary = [
    ("imprezy okolicznościowe Rabka-Zdrój", "100–200", "Średnia", "/imprezy/"),
    ("komunia Rabka-Zdrój restauracja", "50–150", "Niska-Średnia", "/imprezy/"),
    ("chrzciny Rabka-Zdrój", "50–100", "Niska", "/imprezy/"),
    ("urodziny Rabka-Zdrój restauracja", "50–100", "Niska", "/imprezy/"),
    ("bankiet Rabka-Zdrój", "50–150", "Średnia", "/imprezy/"),
    ("impreza firmowa Rabka", "30–80", "Niska", "/imprezy/"),
    ("przyjęcie Rabka-Zdrój", "50–100", "Niska", "/imprezy/"),
    ("catering Rabka-Zdrój", "30–70", "Niska", "/imprezy/"),
]

for row in secondary:
    add_table_row(table3, row)

doc.add_paragraph()

# 3.3 LONG-TAIL
add_heading(doc, "3.3 Long-tail (szybkie wygrane)", level=2)

table4 = doc.add_table(rows=1, cols=3)
table4.style = "Table Grid"
add_table_row(table4, ["Fraza long-tail", "Trudność", "Docelowa strona"], header=True)

longtail = [
    ("góralska chata restauracja Rabka", "Bardzo niska", "/"),
    ("klimatyczna restauracja Rabka-Zdrój", "Bardzo niska", "/"),
    ("tradycyjny obiad Rabka-Zdrój", "Niska", "/menu/"),
    ("pstrąg Rabka-Zdrój restauracja", "Bardzo niska", "/menu/"),
    ("organizacja komunii Rabka-Zdrój 2026", "Bardzo niska", "/imprezy/"),
    ("impreza urodzinowa restauracja Podhale", "Niska", "/imprezy/"),
    ("spotkanie firmowe Rabka-Zdrój", "Niska", "/imprezy/"),
]

for row in longtail:
    add_table_row(table4, row)

# ── 4. IMPLEMENTACJA TECHNICZNA ──────────────────────────────────────────────
doc.add_paragraph()
add_heading(doc, "4. Co zostało wdrożone", level=1)

done_items = [
    ("✅", "robots.txt", "Plik z dyrektywami dla crawlerów + wskazanie sitemap"),
    ("✅", "sitemap.xml", "Mapa wszystkich 5 podstron z priority i lastmod"),
    ("✅", "Meta tagi", "Unikalne title, description, keywords per strona"),
    ("✅", "Open Graph", "og:title, og:description, og:image, og:url dla social media"),
    ("✅", "JSON-LD Schema", "Restaurant schema z adresem, godzinami, oceną, cuisineType"),
    ("✅", "Canonical tags", "Unikalny canonical URL na każdej podstronie"),
    ("✅", "SSR Prerendering", "Statyczny HTML generowany przy buildzie — Google widzi treść bez JS"),
    ("✅", "5 podstron", "/menu/ /imprezy/ /o-nas/ /kontakt/ z unikalnymi H1 i treściami"),
    ("✅", "Routing wouter", "Client-side routing + server-side rendering (prerendering)"),
    ("✅", "Treści SEO", "Każda podstrona ma keyword-rich copy pod lokalne frazy"),
]

table5 = doc.add_table(rows=1, cols=3)
table5.style = "Table Grid"
add_table_row(table5, ["Status", "Element", "Opis"], header=True)
for item in done_items:
    add_table_row(table5, item)

# ── 5. KOLEJNE KROKI ─────────────────────────────────────────────────────────
doc.add_paragraph()
add_heading(doc, "5. Kolejne kroki (do wykonania)", level=1)

next_steps = [
    ("1", "Google Search Console", "Zarejestruj domenę, prześlij sitemap.xml, monitoruj indeksowanie"),
    ("2", "Google Business Profile", "Klient ma już profil — uzupełnij opis, dodaj zdjęcia, odpowiadaj na opinie"),
    ("3", "TripAdvisor listing", "Wysokorankingowy katalog — tworzący widoczność dla turystów"),
    ("4", "Nocowanie.pl", "Lokalny katalog — dużo ruchu z wyszukiwań restauracji w Rabce"),
    ("5", "Resto.com.pl", "Baza restauracji regionu"),
    ("6", "RestauracjePodhale.pl", "Niszowy katalog — niska konkurencja, łatwy link"),
    ("7", "Zdjęcia na stronie", "Optymalizacja nazw plików (IMG_123.jpg → grillowany-oscypek.jpg) + alt teksty"),
    ("8", "Blog / treści", "Artykuły long-tail: 'Gdzie zjeść w Rabce-Zdrój', 'Najlepsze restauracje w Rabce'"),
    ("9", "Opinie Google", "Aktywnie zbieraj nowe opinie — recencja + odpowiedź = sygnał lokalny"),
    ("10", "Local SEO Optimizer", "Pełny audyt lokalny (kolejny krok po wdrożeniu)"),
]

table6 = doc.add_table(rows=1, cols=3)
table6.style = "Table Grid"
add_table_row(table6, ["Priorytet", "Działanie", "Opis"], header=True)
for step in next_steps:
    add_table_row(table6, step)

# ── 6. OCZEKIWANE EFEKTY ─────────────────────────────────────────────────────
doc.add_paragraph()
add_heading(doc, "6. Oczekiwane efekty w czasie", level=1)

table7 = doc.add_table(rows=1, cols=2)
table7.style = "Table Grid"
add_table_row(table7, ["Horyzont czasowy", "Oczekiwany efekt"], header=True)

timeline = [
    ("2 tygodnie", "Google zaindeksuje nowe podstrony — pojawią się w wynikach dla brandowych fraz"),
    ("1 miesiąc", "Pojawienie się w wynikach top 20–30 dla 'restauracja Rabka-Zdrój'"),
    ("2 miesiące", "Wejście do top 10–15 dla głównych fraz; wzrost ruchu z Maps"),
    ("3–6 miesięcy", "Top 5 Google dla głównych fraz lokalnych; widoczność na frazy event-owe"),
    ("Długoterminowo", "Dominacja lokalnego SEO w Rabce-Zdrój — brand recognition"),
]

for row in timeline:
    add_table_row(table7, row)

# ── Footer ───────────────────────────────────────────────────────────────────
doc.add_paragraph()
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Dokument przygotowany przez aisolutions.design · Marzec 2026")
set_font(run, size=9, color=(150, 140, 130), italic=True)

# ── Save ─────────────────────────────────────────────────────────────────────
output_path = os.path.join(os.path.dirname(__file__), "..", "SEO_Strategia_Sielska_Chata.docx")
doc.save(output_path)
print(f"Saved: {os.path.abspath(output_path)}")
